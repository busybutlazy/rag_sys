import json
import os

MAX_PDF_PAGES = int(os.environ.get("PARSER_MAX_PDF_PAGES", "200"))
MAX_EXTRACTED_CHARS = int(os.environ.get("PARSER_MAX_EXTRACTED_CHARS", "500000"))
MAX_JSON_BYTES = int(os.environ.get("PARSER_MAX_JSON_BYTES", "2000000"))
MAX_JSON_DEPTH = int(os.environ.get("PARSER_MAX_JSON_DEPTH", "32"))
MAX_DOCX_PARAGRAPHS = int(os.environ.get("PARSER_MAX_DOCX_PARAGRAPHS", "5000"))


def extract_text(file_path: str, mime_type: str) -> str:
    if mime_type == "application/pdf":
        return _extract_pdf(file_path)
    if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return _extract_docx(file_path)
    if mime_type == "application/json":
        if os.path.getsize(file_path) > MAX_JSON_BYTES:
            raise ValueError("JSON file exceeds parser size limit")
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if _json_depth(data) > MAX_JSON_DEPTH:
            raise ValueError("JSON nesting exceeds parser depth limit")
        return _limit_text(json.dumps(data, ensure_ascii=False, indent=2))
    # text/plain, text/markdown, text/csv — read as-is
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return _limit_text(f.read())


def _extract_pdf(file_path: str) -> str:
    from pypdf import PdfReader
    reader = PdfReader(file_path)
    if len(reader.pages) > MAX_PDF_PAGES:
        raise ValueError("PDF exceeds page limit")
    parts = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return _limit_text("\n\n".join(parts))


def _extract_docx(file_path: str) -> str:
    from docx import Document
    doc = Document(file_path)
    if len(doc.paragraphs) > MAX_DOCX_PARAGRAPHS:
        raise ValueError("DOCX exceeds paragraph limit")
    return _limit_text("\n\n".join(p.text for p in doc.paragraphs if p.text.strip()))


def _limit_text(text: str) -> str:
    if len(text) > MAX_EXTRACTED_CHARS:
        raise ValueError("Extracted text exceeds character limit")
    return text


def _json_depth(value) -> int:
    if isinstance(value, dict):
        return 1 + max((_json_depth(v) for v in value.values()), default=0)
    if isinstance(value, list):
        return 1 + max((_json_depth(v) for v in value), default=0)
    return 1


def chunk_text(text: str, chunk_size: int = 800, chunk_overlap: int = 100) -> list[str]:
    """Sliding-window character chunker with overlap."""
    text = text.strip()
    if not text:
        return []
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end == len(text):
            break
        start = end - chunk_overlap
    return chunks
